"""
Word Document Extractor
"""
import os
import logging
from typing import Dict
from docx import Document

logger = logging.getLogger(__name__)

class DocExtractor:
    def extract(self, file_path: str) -> Dict[str, any]:
        """Extract text from Word document"""
        result = {
            "file_path": file_path,
            "text": "",
            "paragraphs": [],
            "tables": [],
            "metadata": {}
        }
        
        try:
            doc = Document(file_path)
            
            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    result["paragraphs"].append(para.text)
                    result["text"] += para.text + "\n\n"
            
            # Extract tables
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        row_data.append(cell.text.strip())
                    table_data.append(row_data)
                result["tables"].append(table_data)
            
            # Extract metadata
            core_properties = doc.core_properties
            result["metadata"] = {
                "author": core_properties.author,
                "created": str(core_properties.created) if core_properties.created else None,
                "modified": str(core_properties.modified) if core_properties.modified else None,
                "title": core_properties.title,
            }
            
            logger.info(f"Successfully extracted text from {file_path}")
            
        except Exception as e:
            logger.error(f"Error extracting Word document: {e}")
            result["text"] = f"Error extracting document: {str(e)}"
        
        return result
    
    def extract_to_markdown(self, file_path: str) -> str:
        """Extract Word document and convert to markdown"""
        result = self.extract(file_path)
        
        markdown_content = f"# Document: {os.path.basename(file_path)}\n\n"
        
        # Add metadata if available
        if result["metadata"]:
            markdown_content += "## Metadata\n\n"
            for key, value in result["metadata"].items():
                if value:
                    markdown_content += f"- **{key.title()}**: {value}\n"
            markdown_content += "\n"
        
        # Add content
        markdown_content += "## Content\n\n"
        markdown_content += result["text"]
        
        # Add tables
        if result["tables"]:
            markdown_content += "\n## Tables\n\n"
            for idx, table in enumerate(result["tables"]):
                markdown_content += f"### Table {idx + 1}\n\n"
                markdown_content += self._table_to_markdown(table) + "\n\n"
        
        return markdown_content
    
    def _table_to_markdown(self, table: list) -> str:
        """Convert table data to markdown format"""
        if not table:
            return ""
        
        markdown = ""
        for i, row in enumerate(table):
            markdown += "| " + " | ".join(str(cell) for cell in row) + " |\n"
            if i == 0:  # Add header separator
                markdown += "|" + "|".join([" --- " for _ in row]) + "|\n"
        
        return markdown
